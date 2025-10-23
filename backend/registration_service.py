"""
Registration Service for Campus Features
Handles detailed registration for Prize Challenges, Inter-College Competitions, and College Events
"""
from typing import Dict, Any, List, Optional
from fastapi import UploadFile
import os
import uuid
from datetime import datetime, timezone

async def save_student_id_card(file: UploadFile, user_id: str) -> str:
    """Save uploaded student ID card and return URL"""
    try:
        # Create uploads directory if it doesn't exist
        upload_dir = "/app/uploads/student_ids"
        os.makedirs(upload_dir, exist_ok=True)
        
        # Generate unique filename
        file_extension = os.path.splitext(file.filename)[1]
        filename = f"student_id_{user_id}_{uuid.uuid4()}{file_extension}"
        file_path = os.path.join(upload_dir, filename)
        
        # Save file
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        # Return URL path
        return f"/uploads/student_ids/{filename}"
    except Exception as e:
        print(f"Error saving student ID card: {e}")
        return None

async def validate_registration_data(
    registration_type: str,
    data: Dict[str, Any],
    required_individual_fields: List[str] = None,
    required_group_fields: List[str] = None
) -> Dict[str, Any]:
    """Validate registration data based on type"""
    errors = []
    
    if registration_type == "individual":
        required_fields = required_individual_fields or [
            "full_name", "email", "phone_number", "college", 
            "usn", "semester", "year", "branch"
        ]
        
        for field in required_fields:
            if not data.get(field):
                errors.append(f"{field} is required for individual registration")
    
    elif registration_type == "group":
        required_fields = required_group_fields or [
            "team_name", "team_leader_name", "team_leader_email",
            "team_leader_phone", "team_size", "team_members"
        ]
        
        for field in required_fields:
            if not data.get(field):
                errors.append(f"{field} is required for group registration")
        
        # Validate team size
        team_size = data.get("team_size", 0)
        team_members = data.get("team_members", [])
        
        if team_size < 2:
            errors.append("Team size must be at least 2")
        
        if len(team_members) != team_size - 1:  # -1 for team leader
            errors.append(f"Expected {team_size - 1} team members, got {len(team_members)}")
        
        # Validate each team member
        for idx, member in enumerate(team_members):
            member_required = ["name", "email", "phone", "usn"]
            for field in member_required:
                if not member.get(field):
                    errors.append(f"Team member {idx + 1}: {field} is required")
    
    if errors:
        return {"valid": False, "errors": errors}
    
    return {"valid": True, "errors": []}

async def get_registrations_for_event(
    db, 
    event_id: str, 
    event_type: str,
    filters: Dict[str, Any] = None
) -> List[Dict[str, Any]]:
    """Get registrations for an event with optional filters"""
    collection_map = {
        "college_event": "event_registrations",
        "prize_challenge": "prize_challenge_registrations",
        "inter_college": "inter_college_registrations"
    }
    
    collection_name = collection_map.get(event_type)
    if not collection_name:
        return []
    
    # Build query - map event type to correct ID field name
    event_id_field_map = {
        "college_event": "event_id",
        "prize_challenge": "challenge_id",
        "inter_college": "competition_id"
    }
    id_field = event_id_field_map.get(event_type, "event_id")
    query = {id_field: event_id}
    
    if filters:
        # Add college filter
        if filters.get("college"):
            query["$or"] = [
                {"college": filters["college"]},
                {"user_college": filters["college"]},
                {"campus_name": filters["college"]}
            ]
        
        # Add status filter
        if filters.get("status"):
            query["status"] = filters["status"]
        
        # Add registration type filter
        if filters.get("registration_type"):
            query["registration_type"] = filters["registration_type"]
        
        # Add date range filter
        if filters.get("start_date"):
            query["registration_date"] = {"$gte": filters["start_date"]}
        if filters.get("end_date"):
            if "registration_date" not in query:
                query["registration_date"] = {}
            query["registration_date"]["$lte"] = filters["end_date"]
    
    # Fetch registrations
    registrations = await db[collection_name].find(query).to_list(length=1000)
    
    # For inter_college events, fetch from campus_competition_participations (direct joins)
    if event_type == "inter_college":
        participation_query = {"competition_id": event_id}
        
        # Apply campus filter if provided
        if filters and filters.get("college"):
            participation_query["campus"] = filters["college"]
        
        participations = await db.campus_competition_participations.find(participation_query).to_list(length=1000)
        
        # Convert participations to registration format and enrich with user data
        for participation in participations:
            user_id = participation.get("user_id")
            if user_id:
                user = await db.users.find_one({"id": user_id})
                if user:
                    # Get status - check multiple possible field names
                    status = participation.get("status") or participation.get("registration_status", "approved")
                    if status == "registered":
                        status = "approved"  # Map 'registered' to 'approved' for display
                    
                    # Convert participation to registration format
                    reg_item = {
                        "id": participation.get("id"),
                        "competition_id": event_id,
                        "user_id": user_id,
                        "registration_type": "individual",
                        "status": status,
                        "campus_name": participation.get("campus", user.get("university", "Unknown")),
                        "college": participation.get("campus", user.get("university", "Unknown")),
                        # Fields for display in frontend
                        "full_name": user.get("full_name", ""),
                        "user_name": user.get("full_name", ""),
                        "email": user.get("email", ""),
                        "user_email": user.get("email", ""),
                        "phone": user.get("phone", "") or "",
                        "usn": user.get("student_id", "") or user.get("usn", "") or "",
                        # Also keep admin_ prefixed fields for compatibility
                        "admin_name": user.get("full_name", ""),
                        "admin_email": user.get("email", ""),
                        "admin_phone": user.get("phone", "") or "",
                        "admin_usn": user.get("student_id", "") or user.get("usn", "") or "",
                        "registration_date": participation.get("registered_at") or participation.get("joined_at") or participation.get("created_at"),
                        "created_at": participation.get("registered_at") or participation.get("created_at")
                    }
                    registrations.append(reg_item)
    
    # For prize_challenge events, fetch from prize_challenge_participations (direct joins)
    if event_type == "prize_challenge":
        participation_query = {"challenge_id": event_id}
        
        # Apply status filter if provided
        if filters and filters.get("status"):
            # Map status filter to participation_status
            if filters["status"] == "approved":
                participation_query["participation_status"] = {"$in": ["active", "completed"]}
            elif filters["status"] == "pending":
                participation_query["participation_status"] = "pending"
            elif filters["status"] == "rejected":
                participation_query["participation_status"] = "rejected"
        
        participations = await db.prize_challenge_participations.find(participation_query).to_list(length=1000)
        
        # Convert participations to registration format and enrich with user data
        for participation in participations:
            user_id = participation.get("user_id")
            if user_id:
                user = await db.users.find_one({"id": user_id})
                if user:
                    # Map participation_status to registration status
                    part_status = participation.get("participation_status", "active")
                    if part_status in ["active", "completed"]:
                        status = "approved"
                    elif part_status == "pending":
                        status = "pending"
                    else:
                        status = "rejected"
                    
                    # Apply college filter if provided
                    user_college = user.get("university", "Unknown")
                    if filters and filters.get("college") and user_college != filters["college"]:
                        continue
                    
                    # Convert participation to registration format
                    reg_item = {
                        "id": participation.get("id"),
                        "challenge_id": event_id,
                        "user_id": user_id,
                        "registration_type": "individual",
                        "status": status,
                        "campus_name": user_college,
                        "college": user_college,
                        # Fields for display in frontend
                        "full_name": user.get("full_name", ""),
                        "user_name": user.get("full_name", ""),
                        "email": user.get("email", ""),
                        "user_email": user.get("email", ""),
                        "phone": user.get("phone", "") or "",
                        "usn": user.get("student_id", "") or user.get("usn", "") or "",
                        # Also keep admin_ prefixed fields for compatibility
                        "admin_name": user.get("full_name", ""),
                        "admin_email": user.get("email", ""),
                        "admin_phone": user.get("phone", "") or "",
                        "admin_usn": user.get("student_id", "") or user.get("usn", "") or "",
                        "registration_date": participation.get("joined_at") or participation.get("created_at"),
                        "created_at": participation.get("joined_at") or participation.get("created_at"),
                        # Additional prize challenge info
                        "current_progress": participation.get("current_progress", 0),
                        "progress_percentage": participation.get("progress_percentage", 0),
                        "current_rank": participation.get("current_rank")
                    }
                    registrations.append(reg_item)
    
    # Remove MongoDB _id fields
    for reg in registrations:
        if "_id" in reg:
            del reg["_id"]
    
    return registrations

async def get_college_statistics(registrations: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Calculate college-wise statistics from registrations"""
    college_stats = {}
    
    for reg in registrations:
        college = reg.get("college") or reg.get("user_college") or reg.get("campus_name", "Unknown")
        
        if college not in college_stats:
            college_stats[college] = {
                "total": 0,
                "individual": 0,
                "group": 0,
                "pending": 0,
                "approved": 0,
                "rejected": 0
            }
        
        college_stats[college]["total"] += 1
        
        # Count by type
        reg_type = reg.get("registration_type", "individual")
        if reg_type == "individual":
            college_stats[college]["individual"] += 1
        elif reg_type == "group":
            college_stats[college]["group"] += 1
        
        # Count by status
        status = reg.get("status", "pending")
        if status in college_stats[college]:
            college_stats[college][status] += 1
    
    return college_stats

async def export_registrations_to_csv(registrations: List[Dict[str, Any]], event_name: str) -> str:
    """Export registrations to CSV file"""
    import csv
    import io
    
    if not registrations:
        return None
    
    # Create CSV content
    output = io.StringIO()
    
    # Determine fields based on registration type
    if registrations[0].get("registration_type") == "group":
        fieldnames = [
            "Registration ID", "Date", "Status", "Team Name", "Team Leader", 
            "Leader Email", "Leader Phone", "Leader USN", "Team Size", 
            "College", "Branch", "Semester", "Year"
        ]
    else:
        fieldnames = [
            "Registration ID", "Date", "Status", "Name", "Email", "Phone", 
            "USN", "College", "Branch", "Section", "Semester", "Year"
        ]
    
    writer = csv.DictWriter(output, fieldnames=fieldnames)
    writer.writeheader()
    
    for reg in registrations:
        if reg.get("registration_type") == "group":
            row = {
                "Registration ID": reg.get("id", ""),
                "Date": reg.get("registration_date", ""),
                "Status": reg.get("status", ""),
                "Team Name": reg.get("team_name", ""),
                "Team Leader": reg.get("team_leader_name", ""),
                "Leader Email": reg.get("team_leader_email", ""),
                "Leader Phone": reg.get("team_leader_phone", ""),
                "Leader USN": reg.get("team_leader_usn", ""),
                "Team Size": reg.get("team_size", ""),
                "College": reg.get("college", reg.get("user_college", "")),
                "Branch": reg.get("team_leader_branch", reg.get("branch", "")),
                "Semester": reg.get("team_leader_semester", reg.get("semester", "")),
                "Year": reg.get("team_leader_year", reg.get("year", ""))
            }
        else:
            row = {
                "Registration ID": reg.get("id", ""),
                "Date": reg.get("registration_date", ""),
                "Status": reg.get("status", ""),
                "Name": reg.get("full_name", reg.get("user_name", "")),
                "Email": reg.get("email", reg.get("user_email", "")),
                "Phone": reg.get("phone_number", ""),
                "USN": reg.get("usn", ""),
                "College": reg.get("college", reg.get("user_college", "")),
                "Branch": reg.get("branch", ""),
                "Section": reg.get("section", ""),
                "Semester": reg.get("semester", ""),
                "Year": reg.get("year", "")
            }
        writer.writerow(row)
    
    # Save to file
    filename = f"registrations_{event_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
    filepath = f"/app/exports/{filename}"
    
    os.makedirs("/app/exports", exist_ok=True)
    
    with open(filepath, "w") as f:
        f.write(output.getvalue())
    
    return f"/exports/{filename}"


async def export_registrations_to_excel(registrations: List[Dict[str, Any]], event_name: str) -> str:
    """Export registrations to Excel file"""
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    
    if not registrations:
        return None
    
    # Create workbook and worksheet
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Registrations"
    
    # Header style
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # Determine fields based on registration type
    if registrations[0].get("registration_type") == "group":
        headers = [
            "Registration ID", "Date", "Status", "Team Name", "Team Leader", 
            "Leader Email", "Leader Phone", "Leader USN", "Team Size", 
            "College", "Branch", "Semester", "Year"
        ]
    else:
        headers = [
            "Registration ID", "Date", "Status", "Name", "Email", "Phone", 
            "USN", "College", "Branch", "Section", "Semester", "Year"
        ]
    
    # Write headers
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        cell.alignment = Alignment(horizontal='center')
    
    # Write data
    for row_num, reg in enumerate(registrations, 2):
        if reg.get("registration_type") == "group":
            data = [
                reg.get("id", ""),
                reg.get("registration_date", ""),
                reg.get("status", ""),
                reg.get("team_name", ""),
                reg.get("team_leader_name", ""),
                reg.get("team_leader_email", ""),
                reg.get("team_leader_phone", ""),
                reg.get("team_leader_usn", ""),
                reg.get("team_size", ""),
                reg.get("college", reg.get("user_college", "")),
                reg.get("team_leader_branch", reg.get("branch", "")),
                reg.get("team_leader_semester", reg.get("semester", "")),
                reg.get("team_leader_year", reg.get("year", ""))
            ]
        else:
            data = [
                reg.get("id", ""),
                reg.get("registration_date", ""),
                reg.get("status", ""),
                reg.get("full_name", reg.get("user_name", "")),
                reg.get("email", reg.get("user_email", "")),
                reg.get("phone_number", ""),
                reg.get("usn", ""),
                reg.get("college", reg.get("user_college", "")),
                reg.get("branch", ""),
                reg.get("section", ""),
                reg.get("semester", ""),
                reg.get("year", "")
            ]
        
        for col, value in enumerate(data, 1):
            cell = ws.cell(row=row_num, column=col, value=value)
            cell.border = border
    
    # Auto-adjust column widths
    for column in ws.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column_letter].width = adjusted_width
    
    # Save file
    filename = f"registrations_{event_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    filepath = f"/app/exports/{filename}"
    
    os.makedirs("/app/exports", exist_ok=True)
    wb.save(filepath)
    
    return f"/exports/{filename}"


async def export_registrations_to_pdf(registrations: List[Dict[str, Any]], event_name: str) -> str:
    """Export registrations to PDF file"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter, A4, landscape
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    
    if not registrations:
        return None
    
    filename = f"registrations_{event_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    filepath = f"/app/exports/{filename}"
    
    os.makedirs("/app/exports", exist_ok=True)
    
    # Create PDF document
    doc = SimpleDocTemplate(filepath, pagesize=landscape(A4))
    elements = []
    
    # Styles
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    
    # Add title
    title = Paragraph(f"Event Registrations - {event_name}", title_style)
    elements.append(title)
    elements.append(Spacer(1, 20))
    
    # Determine headers based on registration type
    if registrations[0].get("registration_type") == "group":
        headers = [
            "Reg ID", "Date", "Status", "Team Name", "Leader", 
            "Email", "Phone", "USN", "Size", "College"
        ]
        data = [headers]
        for reg in registrations:
            row = [
                reg.get("id", "")[:8],
                reg.get("registration_date", "")[:10],
                reg.get("status", ""),
                reg.get("team_name", "")[:20],
                reg.get("team_leader_name", "")[:15],
                reg.get("team_leader_email", "")[:25],
                reg.get("team_leader_phone", ""),
                reg.get("team_leader_usn", ""),
                str(reg.get("team_size", "")),
                reg.get("college", reg.get("user_college", ""))[:20]
            ]
            data.append(row)
    else:
        headers = [
            "Reg ID", "Date", "Status", "Name", "Email", "Phone", 
            "USN", "College", "Branch", "Semester"
        ]
        data = [headers]
        for reg in registrations:
            row = [
                reg.get("id", "")[:8],
                reg.get("registration_date", "")[:10],
                reg.get("status", ""),
                reg.get("full_name", reg.get("user_name", ""))[:20],
                reg.get("email", reg.get("user_email", ""))[:25],
                reg.get("phone_number", ""),
                reg.get("usn", ""),
                reg.get("college", reg.get("user_college", ""))[:20],
                reg.get("branch", "")[:10],
                reg.get("semester", "")
            ]
            data.append(row)
    
    # Create table
    table = Table(data)
    
    # Table style
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 8),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    elements.append(table)
    
    # Add summary
    elements.append(Spacer(1, 20))
    summary = Paragraph(f"Total Registrations: {len(registrations)}", styles['Normal'])
    elements.append(summary)
    
    # Build PDF
    doc.build(elements)
    
    return f"/exports/{filename}"


async def export_registrations_to_docx(registrations: List[Dict[str, Any]], event_name: str) -> str:
    """Export registrations to DOCX file"""
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    if not registrations:
        return None
    
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Event Registrations - {event_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary
    doc.add_paragraph(f'Total Registrations: {len(registrations)}')
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph()
    
    # Determine columns based on registration type
    if registrations[0].get("registration_type") == "group":
        headers = [
            "Registration ID", "Date", "Status", "Team Name", "Team Leader", 
            "Email", "Phone", "USN", "Team Size", "College"
        ]
    else:
        headers = [
            "Registration ID", "Date", "Status", "Name", "Email", "Phone", 
            "USN", "College", "Branch", "Semester"
        ]
    
    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Add data rows
    for reg in registrations:
        row_cells = table.add_row().cells
        
        if reg.get("registration_type") == "group":
            data = [
                reg.get("id", ""),
                reg.get("registration_date", ""),
                reg.get("status", ""),
                reg.get("team_name", ""),
                reg.get("team_leader_name", ""),
                reg.get("team_leader_email", ""),
                reg.get("team_leader_phone", ""),
                reg.get("team_leader_usn", ""),
                str(reg.get("team_size", "")),
                reg.get("college", reg.get("user_college", ""))
            ]
        else:
            data = [
                reg.get("id", ""),
                reg.get("registration_date", ""),
                reg.get("status", ""),
                reg.get("full_name", reg.get("user_name", "")),
                reg.get("email", reg.get("user_email", "")),
                reg.get("phone_number", ""),
                reg.get("usn", ""),
                reg.get("college", reg.get("user_college", "")),
                reg.get("branch", ""),
                reg.get("semester", "")
            ]
        
        for i, value in enumerate(data):
            row_cells[i].text = str(value)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Save file
    filename = f"registrations_{event_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = f"/app/exports/{filename}"
    
    os.makedirs("/app/exports", exist_ok=True)
    doc.save(filepath)
    
    return f"/exports/{filename}"
